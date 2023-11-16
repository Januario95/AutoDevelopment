#!/usr/bin/env python3
# -*- encoding: utf-8 -*-
"""
Created on Thur May 04, 2023  09:55:41

@author: a248433
"""


import os

X = lambda s: os.system(s)
# X('python -m venv env && env\\Scripts\\activate')
# X('python -m pip install -r requirements.txt')
# os.system('python -m pip install python-docx')
# os.system('python -m pip install docx2pdf')


import os
import time
import locale
import random
from docx import Document
from docx2pdf import convert
from datetime import datetime
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH

# print(dir(WD_SECTION))

locale.setlocale(locale.LC_ALL, 'pt_pt.UTF-8')

class LetterGeneration:
	def __init__(self, company_name, company_address,
	      		 company_phone_nr, account_name, account_nr,
				 branch_code, branch_name, NIB, IBAN,
				 SWIFT, currency, team_leader, manager):
		self.company_name = company_name
		self.company_address = company_address
		self.company_phone_nr = company_phone_nr
		self.account_name = account_name
		self.account_nr = account_nr
		self.branch_code = branch_code
		self.branch_name = branch_name
		self.NIB = NIB
		self.IBAN = IBAN
		self.SWIFT = SWIFT
		self.currency = currency
		self.team_leader = team_leader
		self.manager = manager
		self.document = Document()

	def save(self, filename):
		folder = 'Generate-Files'
		if not os.path.exists(folder):
			os.mkdir('Generate-Files')
		self.document.save(f'{folder}/{filename}.docx')
		time.sleep(1)
		convert(f'{folder}/{filename}.docx')
		time.sleep(1)

		os.system(f'start Generate-Files/{filename}.docx')

	def create_new_section(self, paragraph_type, column_width='2'):
		new_section = self.document.add_section(paragraph_type)
		sectPr = new_section._sectPr
		cols = sectPr.xpath('./w:cols')[0]
		cols.set(qn('w:num'), column_width)

	def generate(self, filename):
		font_styles = self.document.styles
		font_charstyle = font_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
		font_object = font_charstyle.font

		section1 = self.document.sections[0]
		# sectPr = section1._sectPr
		# cols = sectPr.xpath('./w:cols')[0]
		# cols.set(qn('w:num'),'2')

		# paragraph = self.document.add_paragraph()
		# run = paragraph.add_run()
		# run.add_picture('SB_logo_header.png')
		# paragraph.paragraph_format.left_indent = Cm(16.4)

		header = section1.header

		paragraph = header.add_paragraph()
		paragraph.add_run()
		run = paragraph.add_run()
		run.add_picture('SB_logo_header.png')
		# paragraph.paragraph_format.left_indent = Cm(16.4)
		paragraph.paragraph_format.line_spacing = Inches(0)
		paragraph.paragraph_format.space_before = Inches(0)
		paragraph.paragraph_format.space_after = Inches(0)
		paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

		self.create_new_section(WD_SECTION.CONTINUOUS)

		paragraph = self.document.add_paragraph('N/REF/DGE/IS/12/2022/1335\n')
		date=datetime.today()
		month = date.strftime('%B').title()
		date = date.strftime(f'Maputo, %d de {month} de %Y')
		run = paragraph.add_run(date)
		run.add_break(WD_BREAK.LINE)
		paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

		section = self.document.sections[0]
		sectPr = section._sectPr
		cols = sectPr.xpath('./w:cols')[0]
		cols.set(qn('w:num'),'2')

		paragraph = self.document.add_paragraph(f'{self.company_name}\n')
		paragraph.add_run(f'{self.company_address}\n')
		paragraph.add_run(f'Tel: {self.company_phone_nr}')
		paragraph.paragraph_format.line_spacing = Inches(0.25)
		paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
		
		self.create_new_section(WD_SECTION.NEW_COLUMN)

		paragraph = self.document.add_paragraph()
		paragraph.add_run('\n\n\n')
		run = paragraph.add_run('Endereço sede: ')
		run.bold = True
		paragraph.add_run('Standard Bank Sede\n')
		paragraph.add_run('Av. 10 de Novembro, nº 420\n')
		paragraph.add_run('3º andar c. postal 1119\n')
		paragraph.add_run('Maputo')
		paragraph.paragraph_format.line_spacing = Inches(0.25)
		paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

		self.create_new_section(WD_SECTION.CONTINUOUS, 
			  					column_width='1')

		paragraph = self.document.add_paragraph()
		self.bold_text(paragraph, 'ASSUNTO: ', bold_text=True)
		run = paragraph.add_run('Confirmação de Dados Bancários')
		run.underline = True
		paragraph = self.document.add_paragraph('Exmos Senhores,\n')
		# paragraph.paragraph_format.space_before = Inches(0.15)
		# paragraph.paragraph_format.space_after = Inches(0.2)
		paragraph.add_run('Servimo-nos da presente para fornecer os seguintes dados bancários da vossa Organização:')
		paragraph.paragraph_format.line_spacing = Inches(0.25)

		paragraph = self.document.add_paragraph()
		self.bold_text(paragraph, 'Nome da conta:           ')
		paragraph.add_run(f'{self.account_name}\n')
		
		self.bold_text(paragraph, f'Conta:                          ')
		paragraph.add_run(f'{self.account_nr}\n')
		
		self.bold_text(paragraph, 'Código do Balcão:      ')
		paragraph.add_run(f'{self.branch_code}\n')
		
		self.bold_text(paragraph, 'Balcão:                         ')
		paragraph.add_run(f'{self.branch_name}\n')
		
		self.bold_text(paragraph, 'NIB:                             ')
		paragraph.add_run(f'{self.NIB}\n')
		
		self.bold_text(paragraph, 'IBAN:                          ')
		paragraph.add_run(f'{self.IBAN}\n')
		
		self.bold_text(paragraph, 'SWIFT:                       ')
		paragraph.add_run(f'{self.SWIFT}\n')
		
		self.bold_text(paragraph, 'Moeda:                        ')
		paragraph.add_run(f'{self.currency}')
		paragraph.paragraph_format.line_spacing = Inches(0.18)
		paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

		paragraph = self.document.add_paragraph('\nSem mais de momento, subscrevemo-nos com elevada estima e consideração\n')
		paragraph.add_run('De V. Excias,\n')
		paragraph.add_run('Atenciosamente,\n')
		paragraph.paragraph_format.line_spacing = Inches(0.26)
		
		paragraph = self.document.add_paragraph(f'{self.team_leader}       \t\t\t\t{self.manager}\n')
		paragraph.add_run('Team leader Onboarding   \t\t\tManager Onboarding & CEO')
		paragraph.paragraph_format.line_spacing = Inches(0.25)

		# paragraph = self.document.add_paragraph()
		paragraph.add_run()
		run = paragraph.add_run()
		run.add_picture('Stamp.png')
		paragraph.paragraph_format.space_after = Inches(0)

		self.set_font_family()

		# paragraph = self.document.add_paragraph('\n\n')
		# self.bold_text(paragraph, 'Standard Bank Sede', color=RGBColor(0x42, 0x24, 0xE9))
		# self.bold_text(paragraph, ', Avenida 10 de Novembro nº 420/ Caixa Postal 1119/ Maputo\n', bold_text=False, color=RGBColor(0x42, 0x24, 0xE9))
		# self.bold_text(paragraph, 'Tel: +258 21501000/ +25821501100/ 25821501200/ ', bold_text=False, color=RGBColor(0x42, 0x24, 0xE9))
		# self.bold_text(paragraph, 'standardbank.co.mz', color=RGBColor(0x42, 0x24, 0xE9))
		# for run in paragraph.runs:
		# 	run.font.size = Pt(8)
		# 	run.font.name = 'Calibri'

		# paragraph = self.document.add_paragraph()
		# run = paragraph.add_run()
		# run.add_picture('SB_footer.png')
		# paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

		# self.set_margins()
		# self.save(filename)
		
		section = self.document.add_section(WD_SECTION.CONTINUOUS)
		sectPr = section._sectPr
		cols = sectPr.xpath('./w:cols')[0]
		cols.set(qn('w:num'),'1')

		footer = section.footer
		# print(dir(footer))

		paragraph = footer.add_paragraph('\n\n')
		self.bold_text(paragraph, 'Standard Bank Sede', color=RGBColor(40, 98, 128))
		self.bold_text(paragraph, ', Avenida 10 de Novembro nº 420/ Caixa Postal 1119/ Maputo\n', bold_text=False, color=RGBColor(40, 98, 128))
		self.bold_text(paragraph, 'Tel: +258 21501000/ +25821501100/ 25821501200/ ', bold_text=False, color=RGBColor(40, 98, 128))
		self.bold_text(paragraph, 'standardbank.co.mz', color=RGBColor(40, 98, 128))
		for run in paragraph.runs:
			run.font.size = Pt(8)
			run.font.name = 'Calibri'

		paragraph = footer.add_paragraph()
		paragraph.add_run()
		run = paragraph.add_run()
		run.add_picture('SB_footer.png')


		self.set_margins()
		self.save(filename)

	def set_font_family(self):
		for paragraph in self.document.paragraphs:
			paragraph.style = self.document.styles['Normal']
			for run in paragraph.runs:
				run.font.size = Pt(11)
				run.font.name = 'Times New Roman'

	def set_margins(self):
		sections = self.document.sections
		for section in sections:
			section.left_margin = Cm(2)
			section.right_margin = Cm(2)

	
		
	
		
	def bold_text(self, paragraph, text, bold_text=True, color=RGBColor(0, 0, 0)):
		run = paragraph.add_run(text)
		if bold_text:
			run.bold = True
		font = run.font
		font.color.rgb = color
		# font.color = color

# print(help(RGBColor))
generator = LetterGeneration(
	'ECO WEATHER IT & SOLAR SOLUTION LDA',
	'B. FUMENTO AV. PATRICE LUMUMBA',
	'258868894480', 'ECO WEATHER IT & SOLAR SOLUTION LDA',
	'1179030441005', '117', 'S17 AGENCIA DA MATOLA',
	'000301170903044100587', 'MZ59000301170903044100587',
	'SBICMZMX', 'MZN', 'Ivone Faquir', 'Nacima Khan'
)
generator.generate(filename='demo')



def generate_letter(date=datetime.today(), filename='demo'):
	document = Document()
	font_styles = document.styles
	font_charstyle = font_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
	font_object = font_charstyle.font

	section1 = document.sections[0]
	sectPr = section1._sectPr
	cols = sectPr.xpath('./w:cols')[0]
	cols.set(qn('w:num'),'2')

	paragraph = document.add_paragraph()
	run = paragraph.add_run()
	run.add_picture('SB_logo.jpg')
	paragraph.paragraph_format.left_indent = Cm(14)
	# run.add_break(WD_BREAK.LINE)

	# document.add_page_break()

	# new_section = document.add_section(WD_SECTION.NEW_PAGE)
	new_section = document.add_section(WD_SECTION.CONTINUOUS)
	# new_section = document.add_section(WD_SECTION.NEW_COLUMN)
	sectPr = new_section._sectPr
	cols = sectPr.xpath('./w:cols')[0]
	cols.set(qn('w:num'),'2')

	paragraph = document.add_paragraph('N/REF/DGE/IS/12/2022/1335\n')
	month = date.strftime('%B').title()
	date = date.strftime(f'Maputo, %d de {month} de %Y')
	run = paragraph.add_run(date)
	run.add_break(WD_BREAK.LINE)
	paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
	# paragraph.paragraph_format.left_indent = Cm(14)

	section = document.sections[0]
	sectPr = section._sectPr
	cols = sectPr.xpath('./w:cols')[0]
	cols.set(qn('w:num'),'2')

	paragraph = document.add_paragraph('ECO WEATHER IT & SOLAR SOLUTION LDA\n')
	paragraph.add_run('B. FUMENTO AV. PATRICE LUMUMBA\n')
	paragraph.add_run('Tel: 258868894480')
	paragraph.paragraph_format.line_spacing = Inches(0.25)
	paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

	section = document.add_section(WD_SECTION.NEW_COLUMN)
	sectPr = section._sectPr
	cols = sectPr.xpath('./w:cols')[0]
	cols.set(qn('w:num'),'2')
	
	paragraph = document.add_paragraph()
	paragraph.add_run('\n\n\n')
	paragraph.add_run('Endereço sede: Standard Bank Sede\n')
	paragraph.add_run('Av. 10 de Novembro, nº 420\n')
	paragraph.add_run('3º andar c. postal 1119\n')
	paragraph.add_run('Maputo')
	paragraph.paragraph_format.line_spacing = Inches(0.25)
	paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


	section = document.add_section(WD_SECTION.CONTINUOUS)
	sectPr = section._sectPr
	cols = sectPr.xpath('./w:cols')[0]
	cols.set(qn('w:num'),'1')

	paragraph = document.add_paragraph('ASSUNTO: Confirmação de Dados Bancários')

	paragraph = document.add_paragraph('\n\nExmos Senhores,\n')
	paragraph.add_run('\nServimo-nos da presente para fornecer os seguintes dados bancários da vossa Organização:')
	
	paragraph = document.add_paragraph('Nome da conta:          ECO WEATHER IT & SOLAR SOLUTION LDA\n')
	paragraph.add_run('Conta:                         1179030441005\n')
	paragraph.add_run('Código do Balcão:      117\n')
	paragraph.add_run('Balcão:                       S17 AGENCIA DA MATOLA\n')
	paragraph.add_run('NIB:                            000301170903044100587\n')
	paragraph.add_run('IBAN:                          MZ59000301170903044100587\n')
	paragraph.add_run('SWIFT:                       SBICMZMX\n')
	paragraph.add_run('Moeda:                       MZN\n')
	paragraph.paragraph_format.line_spacing = Inches(0.2)
	paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

	paragraph = document.add_paragraph('\nSem mais de momento, subscrevemo-nos com elevada estima e consideração\n\n')
	paragraph.add_run('De V. Excias\n\n')
	paragraph.add_run('Atenciosamente,\n\n\n\n')
	paragraph.add_run('Ivone Faquir       \t\t\t\tNacima Khan\n')
	paragraph.add_run('Team leader Onboarding   \t\t\tManager Onboarding & CEO')
	paragraph.paragraph_format.line_spacing = Inches(0.25)

	# for account in accounts:
	# 	paragraph = document.add_paragraph(f'Conta nº  {account}')
	# 	paragraph.paragraph_format.space_after = Pt(0)
	# 	paragraph.paragraph_format.space_after = Pt(0)


	sections = document.sections
	for section in sections:
		section.left_margin = Cm(3)
		section.right_margin = Cm(3)

	for paragraph in document.paragraphs:
		paragraph.style = document.styles['Normal']
		for run in paragraph.runs:
			run.font.size = Pt(10)
			run.font.name = 'Arial'

	folder = 'Generate-Files'
	if not os.path.exists(folder):
		os.mkdir('Generate-Files')
	document.save(f'{folder}/{filename}.docx')
	time.sleep(1)
	# convert(f'{folder}/{filename}.docx')
	time.sleep(1)

	os.system(f'start Generate-Files/{filename}.docx')



# if __name__=='__main__':
	# accounts = [generate_acc_nr() for _ in range(3)]
	# generate_letter()


