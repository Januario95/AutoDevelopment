#!/usr/bin/env python3
# -*- encoding: utf-8 -*-
"""
Created on Thur May 04, 2023  09:55:41

@author: a248433
"""


import os

X = lambda s: os.system(s)
# X('python -m venv env && env\\Scripts\\activate')
# X('python -m pip install -r requirements.txt')
# os.system('python -m pip install python-docx')
# os.system('python -m pip install docx2pdf')


import os
import time
import shutil
import locale
import random
from docx import Document
from docx2pdf import convert
from datetime import datetime
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH

from _helpers import (
	progress_bar,
)
from _helpers import (
	ClientesSociedadePrenchido,
	ContasSociedadeConstituidaPr, ClientCIFs,
)

locale.setlocale(locale.LC_ALL, 'pt_pt.UTF-8')

class LetterGeneration:
	def __init__(self, company_name=None, company_address=None,
	      		 company_phone_nr=None, account_name=None, account_nr=None,
				 branch_code=None, branch_name=None, NIB=None, IBAN=None,
				 SWIFT=None, currency=None, team_leader=None, manager=None):
		self.company_name = company_name
		self.company_address = company_address
		self.company_phone_nr = company_phone_nr
		self.account_name = account_name
		self.account_nr = account_nr
		self.branch_code = branch_code
		self.branch_name = branch_name
		self.NIB = NIB
		self.IBAN = IBAN
		self.SWIFT = SWIFT
		self.currency = currency
		self.team_leader = team_leader
		self.manager = manager
		self.document = Document()

	def save(self, foldername):
		for char in '!"#$%&\'()*+,-./:;<=>?@[\\]^`{|}~':
			foldername = foldername.replace(char, '')
		filename = foldername.replace(' ', '_')

		# os.chdir(os.getcwd() + f'\\Clientes\\approved\\{filename}')
		# os.chdir(f'C:/Users/a248433/Documents/SB Mozambique/EDO/Robotics/Bots/Automatização do Processo Abertura de contas Corporate/Clientes/approved/{filename}')
		progress_bar(iterations=250)
		self.document.save(f'./{foldername}/{filename}.docx')
		convert(f'./{foldername}/{filename}.docx')
		print(f'FOLDER NAME: {foldername}')
		print(os.getcwd().upper())
		os.system(f'start ./"{foldername}"/{filename}.docx')
		time.sleep(3)
		# shutil.move(f'./{foldername}', '../done')
		# time.sleep()
		try:
			shutil.move(f'./{foldername}', '../done')
			# shutil.move(os.path.join(os.getcwd() + f'/{filename}'),
			# 						'../done')
		except Exception as e:
			print(e.args)
		time.sleep(1)
		# os.unlink(f'./{foldername}/data.xlsx')
		shutil.rmtree(f'/{foldername}', ignore_errors=True)
		# try:
		# 	# os.system(f'rmdir /s ./{foldername!r}')
		# 	shutil.rmtree(f'/{foldername}', ignore_errors=True)
		# 	# os.rmdir(foldername)
		# except Exception as e:
		# 	print(e.args)

	def create_new_section(self, paragraph_type, column_width='2'):
		new_section = self.document.add_section(paragraph_type)
		sectPr = new_section._sectPr
		cols = sectPr.xpath('./w:cols')[0]
		cols.set(qn('w:num'), column_width)

	def generate(self, filename):
		font_styles = self.document.styles
		font_charstyle = font_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
		font_object = font_charstyle.font

		section1 = self.document.sections[0]
		# sectPr = section1._sectPr
		# cols = sectPr.xpath('./w:cols')[0]
		# cols.set(qn('w:num'),'2')

		# paragraph = self.document.add_paragraph()
		# run = paragraph.add_run()
		# run.add_picture('SB_logo_header.png')
		# paragraph.paragraph_format.left_indent = Cm(16.4)

		header = section1.header

		paragraph = header.add_paragraph()
		paragraph.add_run()
		run = paragraph.add_run()
		run.add_picture('C:/Users/a248433/Documents/SB Mozambique/EDO/Robotics/Bots/Automatização do Processo Abertura de contas Corporate/SB_logo_header.png')
		# paragraph.paragraph_format.left_indent = Cm(16.4)
		paragraph.paragraph_format.line_spacing = Inches(0)
		paragraph.paragraph_format.space_before = Inches(0)
		paragraph.paragraph_format.space_after = Inches(0)
		paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

		self.create_new_section(WD_SECTION.CONTINUOUS)

		paragraph = self.document.add_paragraph('N/REF/DGE/IS/12/2022/1335\n')
		date=datetime.today()
		month = date.strftime('%B').title()
		date = date.strftime(f'Maputo, %d de {month} de %Y')
		run = paragraph.add_run(date)
		run.add_break(WD_BREAK.LINE)
		paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

		section = self.document.sections[0]
		sectPr = section._sectPr
		cols = sectPr.xpath('./w:cols')[0]
		cols.set(qn('w:num'),'2')

		paragraph = self.document.add_paragraph(f'{self.company_name}\n')
		paragraph.add_run(f'{self.company_address}\n')
		paragraph.add_run(f'Tel: {self.company_phone_nr}')
		paragraph.paragraph_format.line_spacing = Inches(0.25)
		paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
		
		self.create_new_section(WD_SECTION.NEW_COLUMN)

		paragraph = self.document.add_paragraph()
		paragraph.add_run('\n\n\n')
		run = paragraph.add_run('Endereço sede: ')
		run.bold = True
		paragraph.add_run('Standard Bank Sede\n')
		paragraph.add_run('Av. 10 de Novembro, nº 420\n')
		paragraph.add_run('3º andar c. postal 1119\n')
		paragraph.add_run('Maputo')
		paragraph.paragraph_format.line_spacing = Inches(0.25)
		paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

		self.create_new_section(WD_SECTION.CONTINUOUS, 
			  					column_width='1')

		paragraph = self.document.add_paragraph()
		self.bold_text(paragraph, 'ASSUNTO: ', bold_text=True)
		run = paragraph.add_run('Confirmação de Dados Bancários')
		run.underline = True
		paragraph = self.document.add_paragraph('Exmos Senhores,\n')
		# paragraph.paragraph_format.space_before = Inches(0.15)
		# paragraph.paragraph_format.space_after = Inches(0.2)
		paragraph.add_run('Servimo-nos da presente para fornecer os seguintes dados bancários da vossa Organização:')
		paragraph.paragraph_format.line_spacing = Inches(0.25)

		paragraph = self.document.add_paragraph()
		self.bold_text(paragraph, 'Nome da conta:           ')
		paragraph.add_run(f'{self.account_name}\n')
		
		self.bold_text(paragraph, f'Conta:                          ')
		paragraph.add_run(f'{self.account_nr}\n')
		
		self.bold_text(paragraph, 'Código do Balcão:      ')
		paragraph.add_run(f'{self.branch_code}\n')
		
		self.bold_text(paragraph, 'Balcão:                         ')
		paragraph.add_run(f'{self.branch_name}\n')
		
		self.bold_text(paragraph, 'NIB:                             ')
		paragraph.add_run(f'{self.NIB}\n')
		
		self.bold_text(paragraph, 'IBAN:                          ')
		paragraph.add_run(f'{self.IBAN}\n')
		
		self.bold_text(paragraph, 'SWIFT:                       ')
		paragraph.add_run(f'{self.SWIFT}\n')
		
		self.bold_text(paragraph, 'Moeda:                        ')
		paragraph.add_run(f'{self.currency}')
		paragraph.paragraph_format.line_spacing = Inches(0.18)
		paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

		paragraph = self.document.add_paragraph('\nSem mais de momento, subscrevemo-nos com elevada estima e consideração\n')
		paragraph.add_run('De V. Excias,\n')
		paragraph.add_run('Atenciosamente,\n')
		paragraph.paragraph_format.line_spacing = Inches(0.26)
		
		paragraph = self.document.add_paragraph(f'{self.team_leader}       \t\t\t\t{self.manager}\n')
		paragraph.add_run('Team leader Onboarding   \t\t\tManager Onboarding & CEO')
		paragraph.paragraph_format.line_spacing = Inches(0.25)

		# paragraph = self.document.add_paragraph()
		paragraph.add_run()
		run = paragraph.add_run()
		run.add_picture('C:/Users/a248433/Documents/SB Mozambique/EDO/Robotics/Bots/Automatização do Processo Abertura de contas Corporate/Stamp.png')
		paragraph.paragraph_format.space_after = Inches(0)

		self.set_font_family()

		# paragraph = self.document.add_paragraph('\n\n')
		# self.bold_text(paragraph, 'Standard Bank Sede', color=RGBColor(0x42, 0x24, 0xE9))
		# self.bold_text(paragraph, ', Avenida 10 de Novembro nº 420/ Caixa Postal 1119/ Maputo\n', bold_text=False, color=RGBColor(0x42, 0x24, 0xE9))
		# self.bold_text(paragraph, 'Tel: +258 21501000/ +25821501100/ 25821501200/ ', bold_text=False, color=RGBColor(0x42, 0x24, 0xE9))
		# self.bold_text(paragraph, 'standardbank.co.mz', color=RGBColor(0x42, 0x24, 0xE9))
		# for run in paragraph.runs:
		# 	run.font.size = Pt(8)
		# 	run.font.name = 'Calibri'

		# paragraph = self.document.add_paragraph()
		# run = paragraph.add_run()
		# run.add_picture('SB_footer.png')
		# paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

		# self.set_margins()
		# self.save(filename)
		
		section = self.document.add_section(WD_SECTION.CONTINUOUS)
		sectPr = section._sectPr
		cols = sectPr.xpath('./w:cols')[0]
		cols.set(qn('w:num'),'1')

		footer = section.footer
		# print(dir(footer))

		paragraph = footer.add_paragraph('\n\n')
		self.bold_text(paragraph, 'Standard Bank Sede', color=RGBColor(40, 98, 128))
		self.bold_text(paragraph, ', Avenida 10 de Novembro nº 420/ Caixa Postal 1119/ Maputo\n', bold_text=False, color=RGBColor(40, 98, 128))
		self.bold_text(paragraph, 'Tel: +258 21501000/ +25821501100/ 25821501200/ ', bold_text=False, color=RGBColor(40, 98, 128))
		self.bold_text(paragraph, 'standardbank.co.mz', color=RGBColor(40, 98, 128))
		for run in paragraph.runs:
			run.font.size = Pt(8)
			run.font.name = 'Calibri'

		paragraph = footer.add_paragraph()
		paragraph.add_run()
		run = paragraph.add_run()
		run.add_picture('C:/Users/a248433/Documents/SB Mozambique/EDO/Robotics/Bots/Automatização do Processo Abertura de contas Corporate/SB_footer.png')

		self.set_margins()
		# os.chdir(os.getcwd() + '/Clientes/approved/')
		self.save(filename)

	def set_font_family(self):
		for paragraph in self.document.paragraphs:
			paragraph.style = self.document.styles['Normal']
			for run in paragraph.runs:
				run.font.size = Pt(11)
				run.font.name = 'Times New Roman'

	def set_margins(self):
		sections = self.document.sections
		for section in sections:
			section.left_margin = Cm(2)
			section.right_margin = Cm(2)
	
	def bold_text(self, paragraph, text, bold_text=True, color=RGBColor(0, 0, 0)):
		run = paragraph.add_run(text)
		if bold_text:
			run.bold = True
		font = run.font
		font.color.rgb = color
		# font.color = color

# print(help(RGBColor))
# path = os.getcwd() + f'\\Clientes\\approved\\BANCO NACIONAL DE INVESTIMENTO SA\\data.xlsx'
# client_cif_df = ClientCIFs(filename=path)
# client_sociedade = ClientesSociedadePrenchido(filename=path)
# sociedade_constituida = ContasSociedadeConstituidaPr(filename=path)
# generator = LetterGeneration(
# 	sociedade_constituida.get_value('GB # Nome:'),
# 	str(client_sociedade.get_value('GB # Endereco Residencial')),
# 	str(client_sociedade.get_value('Telefone Celular.1')),
# 	client_sociedade.get_value('GB # Nome:'),
# 	str(sociedade_constituida.get_value('Account')),
# 	'117', 'S17 AGENCIA DA MATOLA',
# 	'000301170903044100587', 'MZ59000301170903044100587',
# 	'SBICMZMX', 'MZN', 'Ivone Faquir', 'Nacima Khan'
# )
# generator.generate(filename=client_sociedade.get_value('GB # Nome:'))



